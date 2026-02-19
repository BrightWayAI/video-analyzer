#!/usr/bin/env python3
"""
Video-to-Storyboard Analyzer for Riptoes

Analyzes educational videos and outputs a shot-by-shot breakdown table.
Helps reverse-engineer effective educational videos to inform storyboard creation.

Usage:
    python analyze_video.py input.mp4 -o breakdown.md
    python analyze_video.py "https://youtube.com/watch?v=xxx" -o breakdown.md
"""

import argparse
import base64
import os
import re
import subprocess
import sys
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

# Third-party imports - these will fail gracefully with helpful messages
try:
    import whisper
except ImportError:
    whisper = None

try:
    from scenedetect import detect, ContentDetector, AdaptiveDetector
except ImportError:
    detect = None
    ContentDetector = None
    AdaptiveDetector = None

try:
    import anthropic
except ImportError:
    anthropic = None

try:
    from PIL import Image
except ImportError:
    Image = None

# docx imports are done lazily in generate_docx_output to avoid import issues
Document = None


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class Shot:
    """Represents a single shot/scene in the video."""
    number: int
    start_time: float  # seconds
    end_time: float    # seconds
    frame_path: Optional[str] = None
    visual_description: str = ""
    dialogue: str = ""

    @property
    def timestamp_str(self) -> str:
        """Format timestamp as M:SS-M:SS."""
        return f"{self._format_time(self.start_time)}-{self._format_time(self.end_time)}"

    @staticmethod
    def _format_time(seconds: float) -> str:
        """Convert seconds to M:SS format."""
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes}:{secs:02d}"


@dataclass
class TranscriptSegment:
    """Represents a segment of transcribed audio."""
    start_time: float
    end_time: float
    text: str
    speaker: str = ""  # Will be inferred


@dataclass
class VideoMetadata:
    """Stores comprehensive video metadata for the storyboard header."""
    title: str
    duration: float
    width: int
    height: int
    orientation: str  # "vertical", "horizontal", or "square"
    word_count: int
    image_count: int
    key_details: str = ""
    reading_level: str = ""
    story_structure: str = ""

    @property
    def duration_str(self) -> str:
        """Format duration as M:SS."""
        minutes = int(self.duration // 60)
        seconds = int(self.duration % 60)
        return f"{minutes}:{seconds:02d}"

    @property
    def aspect_ratio(self) -> str:
        """Return aspect ratio description."""
        if self.orientation == "vertical":
            return f"{self.width}x{self.height} (9:16 vertical)"
        elif self.orientation == "horizontal":
            return f"{self.width}x{self.height} (16:9 horizontal)"
        else:
            return f"{self.width}x{self.height} (square)"


# =============================================================================
# Utility Functions
# =============================================================================

def check_dependencies() -> list[str]:
    """Check for required dependencies and return list of missing ones."""
    missing = []

    # Check FFmpeg
    try:
        subprocess.run(["ffmpeg", "-version"], capture_output=True, check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        missing.append("ffmpeg (system package)")

    # Check Python packages
    if whisper is None:
        missing.append("openai-whisper")
    if detect is None:
        missing.append("scenedetect[opencv]")
    if anthropic is None:
        missing.append("anthropic")
    if Image is None:
        missing.append("Pillow")

    # Check for yt-dlp (only needed for URLs)
    try:
        subprocess.run(["yt-dlp", "--version"], capture_output=True, check=True)
    except (subprocess.CalledProcessError, FileNotFoundError):
        missing.append("yt-dlp (for URL support)")

    return missing


def is_url(path: str) -> bool:
    """Check if the input is a URL."""
    return path.startswith(("http://", "https://", "www."))


def get_video_duration(video_path: str) -> float:
    """Get video duration in seconds using FFprobe."""
    result = subprocess.run(
        [
            "ffprobe", "-v", "error",
            "-show_entries", "format=duration",
            "-of", "default=noprint_wrappers=1:nokey=1",
            video_path
        ],
        capture_output=True,
        text=True,
        check=True
    )
    return float(result.stdout.strip())


def get_video_dimensions(video_path: str) -> tuple[int, int, str]:
    """
    Get video dimensions and orientation using FFprobe.

    Returns:
        Tuple of (width, height, orientation) where orientation is
        "vertical", "horizontal", or "square"
    """
    result = subprocess.run(
        [
            "ffprobe", "-v", "error",
            "-select_streams", "v:0",
            "-show_entries", "stream=width,height",
            "-of", "csv=s=x:p=0",
            video_path
        ],
        capture_output=True,
        text=True,
        check=True
    )
    dimensions = result.stdout.strip().split("x")
    width = int(dimensions[0])
    height = int(dimensions[1])

    # Determine orientation
    if height > width:
        orientation = "vertical"
    elif width > height:
        orientation = "horizontal"
    else:
        orientation = "square"

    return width, height, orientation


def get_video_title(video_path: str) -> str:
    """Extract video title from path or metadata."""
    path = Path(video_path)
    return path.stem


def sanitize_filename(name: str) -> str:
    """Sanitize a string to be safe for use as a filename/folder name."""
    # Replace problematic characters with safe alternatives
    replacements = {
        '/': '-',
        '\\': '-',
        ':': '-',
        '*': '',
        '?': '',
        '"': '',
        '<': '',
        '>': '',
        '|': '-',
        '\n': ' ',
        '\r': '',
    }
    for old, new in replacements.items():
        name = name.replace(old, new)
    # Remove leading/trailing whitespace and dots
    name = name.strip().strip('.')
    # Limit length
    if len(name) > 100:
        name = name[:100]
    return name or "untitled"


# =============================================================================
# Video Download (for URLs)
# =============================================================================

def download_video(url: str, output_dir: str, cookies_file: Optional[str] = None, verbose: bool = False) -> tuple[str, str]:
    """
    Download video from URL using yt-dlp.

    Returns:
        Tuple of (video_path, video_title)
    """
    if verbose:
        print(f"Downloading video from: {url}")

    # Build base command with optional cookies
    base_cmd = ["yt-dlp"]
    if cookies_file and os.path.exists(cookies_file):
        base_cmd.extend(["--cookies", cookies_file])
        if verbose:
            print(f"Using cookies from: {cookies_file}")

    # First get the title
    title_result = subprocess.run(
        base_cmd + ["--get-title", url],
        capture_output=True,
        text=True,
        check=True
    )
    title = title_result.stdout.strip()

    # Download the video with a safe filename to avoid FFmpeg issues with special chars
    safe_filename = "video_download.mp4"
    output_path = os.path.join(output_dir, safe_filename)
    subprocess.run(
        base_cmd + [
            "-f", "bestvideo[ext=mp4]+bestaudio[ext=m4a]/best[ext=mp4]/best",
            "-o", output_path,
            "--no-playlist",
            "--merge-output-format", "mp4",
            url
        ],
        check=True,
        capture_output=not verbose
    )

    if os.path.exists(output_path):
        return output_path, title

    # Fallback: find any video file
    for file in os.listdir(output_dir):
        if file.endswith((".mp4", ".mkv", ".webm", ".avi")):
            return os.path.join(output_dir, file), title

    raise RuntimeError("Failed to find downloaded video file")


# =============================================================================
# Frame Extraction
# =============================================================================

def extract_frames_at_interval(
    video_path: str,
    output_dir: str,
    interval: float = 3.0,
    verbose: bool = False
) -> list[tuple[float, str]]:
    """
    Extract frames at regular intervals using FFmpeg.

    Returns:
        List of (timestamp, frame_path) tuples
    """
    if verbose:
        print(f"Extracting frames every {interval} seconds...")

    os.makedirs(output_dir, exist_ok=True)

    # Extract frames using FFmpeg
    subprocess.run(
        [
            "ffmpeg", "-i", video_path,
            "-vf", f"fps=1/{interval}",
            "-frame_pts", "1",
            os.path.join(output_dir, "frame_%04d.jpg"),
            "-y"
        ],
        check=True,
        capture_output=not verbose
    )

    # Collect frame paths with timestamps
    frames = []
    for i, frame_file in enumerate(sorted(os.listdir(output_dir))):
        if frame_file.startswith("frame_") and frame_file.endswith(".jpg"):
            timestamp = i * interval
            frame_path = os.path.join(output_dir, frame_file)
            frames.append((timestamp, frame_path))

    if verbose:
        print(f"Extracted {len(frames)} frames")

    return frames


def extract_frame_at_time(
    video_path: str,
    timestamp: float,
    output_path: str
) -> str:
    """Extract a single frame at a specific timestamp."""
    subprocess.run(
        [
            "ffmpeg", "-ss", str(timestamp),
            "-i", video_path,
            "-frames:v", "1",
            "-q:v", "2",
            output_path,
            "-y"
        ],
        check=True,
        capture_output=True
    )
    return output_path


# =============================================================================
# Scene Detection
# =============================================================================

def detect_scenes(
    video_path: str,
    min_scene_length: float = 1.5,
    verbose: bool = False
) -> list[tuple[float, float]]:
    """
    Detect scene boundaries using PySceneDetect.

    Returns:
        List of (start_time, end_time) tuples in seconds
    """
    if detect is None:
        raise ImportError("scenedetect is required for scene detection")

    if verbose:
        print("Detecting scene boundaries...")

    # Try multiple detection strategies and combine results
    all_boundaries = set()

    # Strategy 1: ContentDetector with very low threshold for motion graphics
    try:
        scene_list = detect(video_path, ContentDetector(threshold=12.0, min_scene_len=10))
        for scene in scene_list:
            all_boundaries.add(scene[0].get_seconds())
            all_boundaries.add(scene[1].get_seconds())
        if verbose:
            print(f"  ContentDetector found {len(scene_list)} scenes")
    except Exception:
        pass

    # Strategy 2: AdaptiveDetector for gradual transitions (more sensitive)
    try:
        scene_list = detect(video_path, AdaptiveDetector(adaptive_threshold=2.0, min_scene_len=10))
        for scene in scene_list:
            all_boundaries.add(scene[0].get_seconds())
            all_boundaries.add(scene[1].get_seconds())
        if verbose:
            print(f"  AdaptiveDetector found {len(scene_list)} scenes")
    except Exception:
        pass

    # Get video duration for fallback
    duration = get_video_duration(video_path)

    # Convert boundaries to sorted list and create scenes
    boundaries = sorted(all_boundaries)

    # If we have too few boundaries, add interval-based ones
    min_expected_scenes = max(5, int(duration / 8))  # At least 1 scene per 8 seconds
    if len(boundaries) < min_expected_scenes + 1:
        if verbose:
            print(f"  Adding interval markers (expected at least {min_expected_scenes} scenes)")
        interval = duration / (min_expected_scenes + 1)
        for i in range(min_expected_scenes + 2):
            boundaries.append(i * interval)
        boundaries = sorted(set(boundaries))

    # Ensure we have start and end
    if not boundaries or boundaries[0] > 0.5:
        boundaries.insert(0, 0.0)
    if boundaries[-1] < duration - 0.5:
        boundaries.append(duration)

    # Create scenes from boundaries, filtering out very short ones
    scenes = []
    for i in range(len(boundaries) - 1):
        start = boundaries[i]
        end = boundaries[i + 1]
        if end - start >= min_scene_length:
            scenes.append((start, end))

    if verbose:
        print(f"Detected {len(scenes)} scenes")

    return scenes


def create_shots_from_scenes(
    scenes: list[tuple[float, float]],
    video_path: str,
    frames_dir: str,
    verbose: bool = False
) -> list[Shot]:
    """Create Shot objects from detected scenes, extracting representative frames."""
    if verbose:
        print("Extracting representative frames for each scene...")

    os.makedirs(frames_dir, exist_ok=True)
    shots = []

    # Get video duration to avoid extracting frames past the end
    duration = get_video_duration(video_path)

    for i, (start, end) in enumerate(scenes):
        # Extract frame from middle of scene, but stay safely within video bounds
        mid_time = (start + end) / 2
        # Clamp to 0.5 seconds before video end to avoid FFmpeg edge issues
        mid_time = min(mid_time, duration - 0.5)
        mid_time = max(mid_time, 0.1)  # Also avoid very start

        frame_path = os.path.join(frames_dir, f"shot_{i+1:04d}.jpg")
        extract_frame_at_time(video_path, mid_time, frame_path)

        shot = Shot(
            number=i + 1,
            start_time=start,
            end_time=end,
            frame_path=frame_path
        )
        shots.append(shot)

    return shots


def create_shots_from_interval(
    frames: list[tuple[float, str]],
    video_duration: float,
    interval: float
) -> list[Shot]:
    """Create Shot objects from interval-extracted frames."""
    shots = []

    for i, (timestamp, frame_path) in enumerate(frames):
        end_time = min(timestamp + interval, video_duration)
        shot = Shot(
            number=i + 1,
            start_time=timestamp,
            end_time=end_time,
            frame_path=frame_path
        )
        shots.append(shot)

    return shots


# =============================================================================
# Audio Transcription
# =============================================================================

def transcribe_audio(
    video_path: str,
    model_size: str = "base",
    verbose: bool = False
) -> list[TranscriptSegment]:
    """
    Transcribe audio using OpenAI Whisper.

    Returns:
        List of TranscriptSegment objects with timestamps
    """
    if whisper is None:
        raise ImportError("openai-whisper is required for transcription")

    if verbose:
        print(f"Transcribing audio using Whisper ({model_size} model)...")

    # Load model
    model = whisper.load_model(model_size)

    # Transcribe with word-level timestamps
    result = model.transcribe(
        video_path,
        word_timestamps=True,
        verbose=verbose
    )

    # Convert to TranscriptSegment objects
    segments = []
    for segment in result["segments"]:
        ts = TranscriptSegment(
            start_time=segment["start"],
            end_time=segment["end"],
            text=segment["text"].strip()
        )
        segments.append(ts)

    if verbose:
        print(f"Transcribed {len(segments)} segments")

    return segments


def infer_speaker(text: str, context: list[TranscriptSegment] = None) -> str:
    """
    Infer speaker type from text content.

    Returns speaker label like "NARRATOR:", "CHILD:", "TEACHER:", etc.
    """
    text_lower = text.lower()

    # Check for common narrator patterns
    narrator_patterns = [
        "today we", "let's learn", "welcome to", "in this video",
        "once upon", "our story", "let me tell", "did you know"
    ]
    if any(pattern in text_lower for pattern in narrator_patterns):
        return "NARRATOR"

    # Check for question patterns (often teachers)
    if "?" in text and any(w in text_lower for w in ["class", "who can", "what do you", "can anyone"]):
        return "TEACHER"

    # Check for child-like speech patterns
    child_patterns = ["i saw", "i found", "look at", "wow", "cool", "mommy", "daddy"]
    if any(pattern in text_lower for pattern in child_patterns):
        return "CHILD"

    # Default to generic speaker
    return "SPEAKER"


# =============================================================================
# Visual Analysis with Claude Vision
# =============================================================================

def analyze_frame_with_claude(
    frame_path: str,
    client: "anthropic.Anthropic",
    orientation: str = "horizontal",
    verbose: bool = False
) -> str:
    """
    Analyze a video frame using Claude Vision API.

    Args:
        frame_path: Path to the frame image
        client: Anthropic client
        orientation: Video orientation ("vertical", "horizontal", or "square")
        verbose: Print progress

    Returns:
        Visual description of the frame
    """
    if verbose:
        print(f"  Analyzing frame: {os.path.basename(frame_path)}")

    # Read and encode image
    with open(frame_path, "rb") as f:
        image_data = base64.standard_b64encode(f.read()).decode("utf-8")

    # Determine media type
    ext = Path(frame_path).suffix.lower()
    media_type = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".gif": "image/gif",
        ".webp": "image/webp"
    }.get(ext, "image/jpeg")

    # Adapt prompt based on video orientation
    if orientation == "vertical":
        format_note = "vertical 9:16 format"
        framing_guidance = """For vertical format, use these shot types if relevant:
- Full body (head to toe, centered)
- Medium shot (waist up)
- Close-up (face/head)
- Extreme close-up (detail)
- High angle / low angle
- Stacked two-shot (figures arranged vertically)
Do NOT use "wide shot" - vertical format shows depth/height, not width."""
    elif orientation == "square":
        format_note = "square 1:1 format"
        framing_guidance = """For square format, subjects are typically centered. Use close-up, medium shot, or full body as appropriate."""
    else:
        format_note = "horizontal 16:9 format"
        framing_guidance = """For horizontal format, standard shot types apply: wide/establishing shot, medium shot, close-up, two-shot, over-the-shoulder, etc."""

    # Create prompt for AI image generation
    prompt = f"""You are writing an image generation prompt for an AI tool to recreate this frame as a still image for an animated video.

This is a {format_note} video frame.

{framing_guidance}

Write a prompt that would generate this exact image. Include:
- Art style (2D animation, 3D render, realistic, cartoon, etc.)
- Setting/environment with specific details
- Characters: age, appearance, clothing, pose, expression, action
- Lighting and mood
- Shot type/framing ONLY if essential to the composition

Keep the prompt under 50 words. Be specific and concrete - avoid vague terms.

Example outputs:
- "2D cartoon style. Bright classroom interior. Young Black girl with pigtails in yellow dress sits at wooden desk, pencil in hand, looking up thoughtfully. Warm natural lighting from windows. Medium shot."
- "Colorful flat animation. Forest clearing at sunset. Friendly brown bear standing upright, waving paw, smiling. Orange and pink sky visible through trees."

Provide ONLY the image prompt, no preamble or explanation."""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=150,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": media_type,
                            "data": image_data,
                        },
                    },
                    {
                        "type": "text",
                        "text": prompt
                    }
                ],
            }
        ],
    )

    return response.content[0].text.strip()


def analyze_all_frames(
    shots: list[Shot],
    orientation: str = "horizontal",
    verbose: bool = False
) -> list[Shot]:
    """Analyze all shot frames using Claude Vision API."""
    if anthropic is None:
        raise ImportError("anthropic package is required for visual analysis")

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY environment variable is required")

    client = anthropic.Anthropic(api_key=api_key)

    if verbose:
        print(f"Analyzing {len(shots)} frames with Claude Vision...")

    for shot in shots:
        if shot.frame_path and os.path.exists(shot.frame_path):
            shot.visual_description = analyze_frame_with_claude(
                shot.frame_path,
                client,
                orientation,
                verbose
            )

    return shots


def analyze_video_content(
    transcript: list[TranscriptSegment],
    shots: list[Shot],
    title: str,
    verbose: bool = False
) -> tuple[str, str, str]:
    """
    Analyze video content to extract key details, reading level, and story structure.

    Returns:
        Tuple of (key_details, reading_level, story_structure)
    """
    if anthropic is None:
        raise ImportError("anthropic package is required for content analysis")

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("ANTHROPIC_API_KEY environment variable is required")

    client = anthropic.Anthropic(api_key=api_key)

    if verbose:
        print("Analyzing video content for metadata...")

    # Combine transcript text
    full_transcript = " ".join([seg.text for seg in transcript])

    # Combine visual descriptions
    visual_summary = " | ".join([shot.visual_description for shot in shots if shot.visual_description])

    prompt = f"""Analyze this educational video content and provide structured metadata.

Video Title: {title}

Transcript:
{full_transcript}

Visual Summary (scene descriptions):
{visual_summary}

Provide your analysis in EXACTLY this format (keep each section brief, 1-2 sentences max):

KEY DETAILS: [Main topic, target audience, educational objectives - be specific and concise]

READING LEVEL: [Grade level like "Grade 1-2" or "Pre-K" based on vocabulary and concepts]

STORY STRUCTURE: [Narrative format like "Problem-Solution", "Sequential Tutorial", "Character Journey", "Question-Answer", "Cause-Effect", etc. with brief explanation]

Respond with ONLY these three sections, nothing else."""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=500,
        messages=[{"role": "user", "content": prompt}]
    )

    response_text = response.content[0].text.strip()

    # Parse the response
    key_details = ""
    reading_level = ""
    story_structure = ""

    for line in response_text.split("\n"):
        line = line.strip()
        if line.startswith("KEY DETAILS:"):
            key_details = line.replace("KEY DETAILS:", "").strip()
        elif line.startswith("READING LEVEL:"):
            reading_level = line.replace("READING LEVEL:", "").strip()
        elif line.startswith("STORY STRUCTURE:"):
            story_structure = line.replace("STORY STRUCTURE:", "").strip()

    return key_details, reading_level, story_structure


# =============================================================================
# Stylistic Fingerprint v3 (Production-Grade Deterministic Classifier — 8 Fields)
# =============================================================================
#
# Architecture:
#   Phase 1 — Feature Counter Pass: scan every shot description once, compute
#             ratios for ~15 feature families.  These ratios are the ONLY
#             inputs to the 8 classifiers (no duplicate keyword scanning).
#   Phase 2 — Priority-ordered hard-rule classifiers for each of the 8 fields.
#   Phase 3 — Soft-score fallback when no hard rule fires.
#
# Design goals:
#   • Deterministic (no LLM calls).
#   • Single value per field, never "Unknown".
#   • Discriminative enough to reverse-engineer visual identity constraints
#     and derive prompting constraints for generating similar storyboards.
# =============================================================================

# ── Feature Family Keyword Lists (module-level constants) ────────────────

_FF_3D_CUES = [
    "3d", "3d render", "3d animated", "3d animation", "3d character",
    "3d model", "3d realistic", "rendered", "cgi", "pixar", "3-d",
    "three-dimensional", "isometric", "blender", "unreal",
    "depth of field", "volumetric", "realistic render", "ray traced",
    "cinematic lighting",
]

_FF_2D_FLAT_CUES = [
    "flat vector", "flat illustration", "flat design", "flat animation",
    "cel shaded", "clean vector", "geometric shapes", "simple shapes",
    "solid background", "cartoon style", "illustrated",
]

_FF_2D_LINEART_CUES = [
    "line art", "stick figure", "black and white", "paper texture",
    "hand-drawn outline", "line drawing", "minimalist line",
    "black line", "simple black", "crumpled paper",
]

_FF_2D_TEXTURED_CUES = [
    "watercolor", "cross-hatching", "crosshatch", "crosshatching",
    "engraving", "victorian", "ink wash", "etched", "woodcut",
    "grainy illustration", "vintage illustration", "oil painting style",
    "victorian illustration", "etching",
]

_FF_2D_BROAD = [
    "2d", "2d animated", "2d animation", "2d character",
    "2d illustrated", "2d cartoon", "2d black",
    "cartoon", "illustrated",
]

_FF_PHOTOREAL_CUES = [
    "realistic", "photographic", "photoreal", "realistic render",
    "shallow depth of field", "realistic 3d", "live action",
    "documentary", "real footage", "archival", "photograph",
    "historical photo",
]

_FF_ABSTRACT_BG_CUES = [
    "white background", "plain background", "paper texture", "parchment",
    "gradient backdrop", "black void", "no background", "floating objects",
    "textured white", "clean background", "solid background",
    "crumpled paper texture",
]

_FF_REAL_ENV_CUES = [
    "room", "kitchen", "office", "street", "house", "bank", "classroom",
    "factory", "battlefield", "manor", "stairs", "table", "bedroom",
    "library", "store", "shop", "apartment", "courtroom", "workshop",
    "café", "restaurant", "neighborhood", "marketplace", "market",
    "village", "interior", "landscape", "nature", "road", "farm",
    "park", "city", "field", "wall", "furniture", "building",
    "indoor", "outdoor", "sky",
]

_FF_DATA_PRESENTATION_CUES = [
    "chart", "graph", "formula", "timeline", "bullet list", "title card",
    "slide", "diagram", "ui element", "ui screen", "phone screen",
    "app screen", "browser", "dashboard", "logo", "seal", "infographic",
    "statistics", "percentage", "data visualization", "bar chart",
    "pie chart", "calculator",
    # Note: "ui" alone removed — matches "suit". "label" removed — matches
    # "labeled" (e.g. "crate labeled DEBT").  Use specific multi-word forms.
]

_FF_CTA_SOCIAL_CUES_REGEX = re.compile(
    r'\b(?:subscribe|comment|follow|share|part\s*[12]|link in bio)\b',
    re.IGNORECASE,
)

_FF_SYMBOLISM_CUES = [
    "symbolic", "metaphor", "allegory", "hourglass", "hook",
    "fishing hook", "trap", "chain", "weight", "anchor", "clock",
    "bomb", "fire", "iceberg", "bridge", "crushing press", "lever",
    "whistle", "smoke alarm", "scale", "balance",
]

_FF_OBJECT_ANALOGY_CUES = [
    "banana", "juice", "coin", "piggy bank", "golden egg", "treasure",
    "chest", "key", "lock", "seed", "tree growing", "jar",
]

_FF_MASCOT_CUES = [
    "monkey", "duck", "bear", "bunny", "rabbit", "fox", "penguin",
    "mascot", "anthropomorphic", "animal character",
]

_FF_HUMAN_CHARACTER_CUES = [
    "character", "person", "boy", "girl", "child", "man", "woman",
    "kid", "figure", "presenter", "gentleman", "lady", "teen",
    "teenager", "protagonist", "maid", "servant", "worker",
    "narrator", "host", "butler", "housekeeper",
]

_FF_TEXT_OVERLAY_CUES = [
    "text overlay", "title card", "text reading", "centered text",
    "large text", "bold text", "typography", "on-screen text",
    "white text", "kinetic text", "text animation",
]

_FF_CINEMATIC_CUES = [
    "depth of field", "cinematic lighting", "dramatic", "moody",
    "volumetric", "bokeh", "shallow depth", "lens flare",
    "dramatic lighting",
]

_FF_MOTION_CUES = [
    "pan", "zoom", "tracking", "dolly", "tilt", "camera move",
    "motion", "fly", "swooping", "rotating", "orbit",
]

_FF_FICTIONAL_WORLD_CUES = [
    "jungle", "banana", "monkey", "duck", "animal", "fantasy",
    "magical", "mythical", "fairy", "enchanted", "kingdom",
    "mascot", "anthropomorphic", "juice", "treasure", "golden egg",
    "spaceship", "alien", "underwater city", "pirate",
]

_FF_STYLIZED_RENDERING_CUES = [
    "cartoon", "stylized", "3d animated style", "3d animated",
    "animated style", "2d animated", "2d illustration",
    "illustrated", "animation style",
]

# Institutional visual markers
_FF_INSTITUTIONAL_CUES = [
    "logo", "seal", "ted-ed", "ted talk", "university", "museum",
    "official", "historical documentary", "government",
    "education seal", "institutional", "formal title", "regulation",
    "department of", "ministry",
]

# Historical narration transcript cues
_FF_HISTORICAL_NARRATION_CUES = [
    "in the 19th century", "during this era", "historians",
    "century", "historical", "in the 18th century", "in the 1800s",
    "in the 1700s", "medieval", "ancient", "in the early 1900s",
    "during the war", "colonial", "empire",
]

# Corporate visual markers
_FF_CORPORATE_CUES = [
    "office", "interview", "negotiation", "salary", "corporate",
    "business meeting", "professional setting", "meeting room",
    "boardroom", "colleague", "promotion", "desk", "laptop",
    "conference",
]

# GenZ slang (word-boundary regex to avoid false positives)
_FF_GENZ_SLANG_REGEX = re.compile(
    r"\b(?:yo|bruh|lit|no cap|lowkey|vibe|slay|fr fr|real talk|"
    r"bro|boss|you know what|don't lose money|bet|sus|cap|"
    r"rizz|goated|bussin|fam)\b",
    re.IGNORECASE,
)

# Child-friendly visual cues
_FF_CHILD_FRIENDLY_CUES = [
    "anthropomorphic", "bright", "playful", "colorful cartoon",
    "friendly", "cute", "adorable", "whimsical", "mascot",
    "big eyes", "soft colors", "rainbow",
]

# Dark / editorial visual cues
_FF_DARK_EDITORIAL_CUES = [
    "moody", "dramatic lighting", "crosshatch", "engraving",
    "somber", "trap", "symbolic", "victorian", "gloomy",
    "ominous", "sinister", "shadowy", "dark palette",
    "cross-hatching", "dark teal", "dark background",
]

# Moral / cautionary transcript cues
_FF_MORAL_CUES = [
    "penalty", "trap", "charge", "fee", "broke", "debt", "cost",
    "exploit", "unfair", "predatory", "warning", "beware",
    "critique", "irony", "satire", "scam", "hidden cost",
]

# Narrative arc transcript cues
_FF_NARRATIVE_ARC_CUES = [
    "once upon", "story", "journey", "adventure", "problem",
    "resolution", "hero", "discovers", "learns", "overcomes",
    "struggles", "finally", "eventually", "at first", "but then",
]

# Pacing qualifier keywords
_FF_PACING_FAST_CUES = [
    "quick cut", "rapid", "flash", "montage", "fast", "whip",
]

# Shot framing keywords (for camera language detection)
_FF_CLOSE_UP_CUES = [
    "close-up", "close up", "extreme close-up", "detail shot",
    "macro", "tight shot",
]
_FF_WIDE_SHOT_CUES = [
    "wide shot", "establishing shot", "wide angle", "full body",
    "landscape shot", "panoramic",
]
_FF_MEDIUM_SHOT_CUES = [
    "medium shot", "waist up", "mid shot", "two-shot",
    "over-the-shoulder",
]


def analyze_stylistic_fingerprint(
    shots: list[Shot],
    transcript: list[TranscriptSegment],
    duration: float,
    orientation: str = "horizontal",
    reading_level: str = "",
) -> dict:
    """
    Produce a Stylistic Fingerprint (v3) for the video.

    Production-grade deterministic classifier with 8 fields:
        1. Rendering Class
        2. World Type
        3. Character Strategy
        4. Narrative Structure        (NEW in v3 — replaces Metaphor Mode)
        5. Visual Abstraction Index   (NEW in v3 — graduated 1-5 scale)
        6. Visual Density
        7. Camera / Editing Language
        8. Tonal Positioning

    Returns a dict with 8 keys.  Every value is a non-empty string.
    Also returns legacy v2 keys for backward compatibility during transition.
    """
    from collections import Counter

    # ================================================================
    # PHASE 1 — Feature Counter Pass
    # ================================================================
    visual_list = [
        s.visual_description.lower()
        for s in shots
        if s.visual_description
    ]
    total = len(visual_list) if visual_list else 1

    full_transcript_lower = " ".join(seg.text.lower() for seg in transcript)
    full_transcript_raw = " ".join(seg.text for seg in transcript)

    # ── Helper: fraction of shots matching ANY keyword ────────────────
    def _shot_ratio(keywords: list[str]) -> float:
        hits = sum(
            1 for v in visual_list
            if any(kw in v for kw in keywords)
        )
        return hits / total

    # ── Helper: per-shot feature vector (True/False per family) ──────
    # This lets us compute co-occurrence cheaply.
    def _shot_hits(keywords: list[str]) -> list[bool]:
        return [any(kw in v for kw in keywords) for v in visual_list]

    # ── Compute all feature ratios in one pass ───────────────────────
    r_3d          = _shot_ratio(_FF_3D_CUES)
    r_2d_flat     = _shot_ratio(_FF_2D_FLAT_CUES)
    r_lineart     = _shot_ratio(_FF_2D_LINEART_CUES)
    r_textured    = _shot_ratio(_FF_2D_TEXTURED_CUES)
    r_2d_broad    = _shot_ratio(_FF_2D_BROAD)
    r_photoreal   = _shot_ratio(_FF_PHOTOREAL_CUES)
    r_abstract_bg = _shot_ratio(_FF_ABSTRACT_BG_CUES)
    r_real_env    = _shot_ratio(_FF_REAL_ENV_CUES)
    r_data_pres   = _shot_ratio(_FF_DATA_PRESENTATION_CUES)
    r_symbolism   = _shot_ratio(_FF_SYMBOLISM_CUES)
    r_analogy     = _shot_ratio(_FF_OBJECT_ANALOGY_CUES)
    r_mascot      = _shot_ratio(_FF_MASCOT_CUES)
    r_human_char  = _shot_ratio(_FF_HUMAN_CHARACTER_CUES)
    r_text_overlay = _shot_ratio(_FF_TEXT_OVERLAY_CUES)
    r_cinematic   = _shot_ratio(_FF_CINEMATIC_CUES)
    r_motion      = _shot_ratio(_FF_MOTION_CUES)
    r_fictional   = _shot_ratio(_FF_FICTIONAL_WORLD_CUES)
    r_stylized    = _shot_ratio(_FF_STYLIZED_RENDERING_CUES)
    r_institutional = _shot_ratio(_FF_INSTITUTIONAL_CUES)
    r_corporate   = _shot_ratio(_FF_CORPORATE_CUES)
    r_child_vis   = _shot_ratio(_FF_CHILD_FRIENDLY_CUES)
    r_dark_vis    = _shot_ratio(_FF_DARK_EDITORIAL_CUES)
    r_close_up    = _shot_ratio(_FF_CLOSE_UP_CUES)
    r_wide_shot   = _shot_ratio(_FF_WIDE_SHOT_CUES)
    r_medium_shot = _shot_ratio(_FF_MEDIUM_SHOT_CUES)

    # Transcript-level boolean flags
    has_cta = bool(_FF_CTA_SOCIAL_CUES_REGEX.search(full_transcript_lower))
    has_genz_slang = bool(_FF_GENZ_SLANG_REGEX.search(full_transcript_lower))
    has_historical = any(
        kw in full_transcript_lower for kw in _FF_HISTORICAL_NARRATION_CUES
    )
    has_moral_framing = any(
        kw in full_transcript_lower for kw in _FF_MORAL_CUES
    )
    has_narrative_arc = any(
        kw in full_transcript_lower for kw in _FF_NARRATIVE_ARC_CUES
    )

    # Timing features
    durations = [
        s.end_time - s.start_time
        for s in shots
        if s.end_time > s.start_time
    ]
    avg_shot_len = (sum(durations) / len(durations)) if durations else 5.0
    is_vertical = (orientation == "vertical")
    is_horizontal = (orientation == "horizontal")
    is_short_form = (duration <= 90)

    # Parse reading level into numeric grade
    grade_num = 99
    if reading_level:
        _rl = reading_level.lower().replace("pre-k", "0").replace("prek", "0")
        _gm = re.search(r'(\d+)', _rl)
        if _gm:
            grade_num = int(_gm.group(1))

    # Co-occurrence: shots that have BOTH abstract bg AND data cues
    _hits_abstract = _shot_hits(_FF_ABSTRACT_BG_CUES)
    _hits_data     = _shot_hits(_FF_DATA_PRESENTATION_CUES)
    _hits_env      = _shot_hits(_FF_REAL_ENV_CUES)
    _hits_3d       = _shot_hits(_FF_3D_CUES)
    _hits_2d       = _shot_hits(_FF_2D_BROAD)
    co_abstract_data = sum(
        1 for a, d in zip(_hits_abstract, _hits_data) if a and d
    ) / total
    co_env_stylized = sum(
        1 for e, s in zip(_hits_env, _shot_hits(_FF_STYLIZED_RENDERING_CUES))
        if e and s
    ) / total

    # ================================================================
    # PHASE 2 — Eight Classifiers
    # ================================================================

    # ══════════════════════════════════════════════════════════════════
    # 1) Rendering Class
    # ══════════════════════════════════════════════════════════════════
    # Priority order with hard thresholds, then soft-score fallback.
    rendering_class = None

    # Mixed Media: significant BOTH 3D and 2D signals, neither dominant
    if (r_3d >= 0.20 and r_2d_broad >= 0.20
            and r_3d < 0.60 and r_2d_broad < 0.60):
        rendering_class = "Mixed Media (2D + 3D)"
    # 3D dominant
    elif r_3d >= 0.40:
        if r_photoreal >= 0.20:
            rendering_class = "Photoreal / Realistic 3D"
        else:
            rendering_class = "Stylized 3D"
    # Minimalist Line Art
    elif r_lineart >= 0.40:
        rendering_class = "Minimalist Line Art 2D"
    # Textured 2D (engraving/watercolor)
    elif r_textured >= 0.30:
        rendering_class = "Textured 2D (engraving/watercolor)"
    # Flat 2D
    elif r_2d_flat >= 0.30:
        rendering_class = "Flat 2D"

    # Soft scoring fallback
    if rendering_class is None:
        _rc_scores = {
            "Stylized 3D": r_3d,
            "Flat 2D": r_2d_flat + r_2d_broad * 0.3,
            "Minimalist Line Art 2D": r_lineart,
            "Textured 2D (engraving/watercolor)": r_textured,
            "Mixed Media (2D + 3D)": min(r_3d, r_2d_broad),
        }
        rendering_class = max(_rc_scores, key=_rc_scores.get)
        if all(v == 0 for v in _rc_scores.values()):
            rendering_class = "Flat 2D"

    # ══════════════════════════════════════════════════════════════════
    # 2) World Type  (background-sensitive, priority order)
    # ══════════════════════════════════════════════════════════════════
    # Priority 1: Abstract Concept Space
    #   High abstract-bg AND low environment AND not data-dominated
    if r_abstract_bg >= 0.60 and r_real_env < 0.20 and r_data_pres < 0.30:
        world_type = "Abstract Concept Space"
    # Priority 2: Data/Presentation Space
    elif r_data_pres >= 0.30:
        world_type = "Data/Presentation Space"
    elif r_abstract_bg >= 0.40 and co_abstract_data >= 0.20:
        world_type = "Data/Presentation Space"
    # Priority 3: Fictional Metaphor Universe
    elif r_fictional >= 0.30:
        world_type = "Fictional Metaphor Universe"
    # Priority 4: Stylized Real-World
    elif r_real_env >= 0.30 and r_stylized >= 0.20:
        world_type = "Stylized Real-World"
    elif r_real_env >= 0.30 and (r_3d >= 0.20 or r_2d_broad >= 0.20):
        world_type = "Stylized Real-World"
    # Priority 5: Real-World Literal
    elif r_photoreal >= 0.20 or (r_real_env >= 0.30 and r_stylized < 0.10):
        world_type = "Real-World Literal"
    else:
        # Fallback
        if r_abstract_bg >= 0.30:
            world_type = "Abstract Concept Space"
        else:
            world_type = "Stylized Real-World"

    # ══════════════════════════════════════════════════════════════════
    # 3) Character Strategy  (identity-aware)
    # ══════════════════════════════════════════════════════════════════
    # ── Identity detection: find DISTINCT characters ──
    _identity_patterns = [
        # Named species/role labels
        r'(?:monkey|duck|bear|bunny|rabbit|fox|penguin)'
        r'\s*(?:mascot|banker|professor|teacher|character)?',
        # Specific identity: adjective + noun combos that recur
        r'(?:young |little |old |teenage |elderly )?'
        r'(?:black |white |brown |blonde )?'
        r'(?:boy|girl|child|kid|man|woman|gentleman|lady|teen|teenager|'
        r'maid|servant|butler|housekeeper)'
        r'\s*(?:with [^,.]+|in [^,.]+)?',
    ]
    char_identities = []
    for si, v in enumerate(visual_list):
        for pat in _identity_patterns:
            for match in re.findall(pat, v):
                label = match.strip()
                if label:
                    char_identities.append((label, si))

    identity_shot_map: dict[str, set[int]] = {}
    for label, si in char_identities:
        identity_shot_map.setdefault(label, set()).add(si)

    # Distinct = label in >= 3 shots
    distinct_characters = {
        label: shots_set
        for label, shots_set in identity_shot_map.items()
        if len(shots_set) >= 3
    }

    # Also check transcript for proper names (≥2 occurrences)
    _proper_names = re.findall(
        r'(?<!\. )(?<!\.\n)\b([A-Z][a-z]{2,})\b', full_transcript_raw
    )
    _name_counts = Counter(_proper_names)
    _stop_words = {
        "the", "and", "but", "this", "that", "here", "there",
        "when", "what", "how", "then", "from", "with", "for",
        "not", "are", "was", "has", "had", "have",
    }
    for name, cnt in _name_counts.items():
        if cnt >= 2 and name.lower() not in _stop_words:
            name_shots = {
                i for i, v in enumerate(visual_list) if name.lower() in v
            }
            if len(name_shots) >= 2:
                distinct_characters[name.lower()] = name_shots

    num_distinct = len(distinct_characters)
    max_char_coverage = max(
        (len(s) / total for s in distinct_characters.values()),
        default=0.0,
    )

    # Object-only indicators
    _object_only_kw = [
        "chart", "graph", "coin", "icon", "hourglass", "hook",
        "scale", "piggy bank", "text overlay", "title card",
        "logo", "symbol", "object",
    ]
    r_objects = _shot_ratio(_object_only_kw)

    # Classification (priority order)
    if r_human_char < 0.30 and r_mascot < 0.20 and (r_objects >= 0.30 or num_distinct == 0):
        character_strategy = "None (objects only)"
    elif r_mascot >= 0.40:
        character_strategy = "Mascot-Led (non-human recurring character)"
    elif max_char_coverage >= 0.50 and num_distinct <= 2:
        if has_narrative_arc:
            character_strategy = "Single Protagonist Arc"
        else:
            character_strategy = "Single Narrator/Host Character"
    elif max_char_coverage >= 0.40 and num_distinct <= 2:
        character_strategy = "Single Narrator/Host Character"
    elif num_distinct >= 3:
        chars_with_coverage = [
            label for label, s in distinct_characters.items()
            if len(s) / total >= 0.20
        ]
        if len(chars_with_coverage) >= 3:
            character_strategy = "Ensemble Cast"
        else:
            character_strategy = "Single Narrator/Host Character"
    elif r_human_char >= 0.30:
        character_strategy = "Single Narrator/Host Character"
    else:
        character_strategy = "None (objects only)"

    # ══════════════════════════════════════════════════════════════════
    # 4) Narrative Structure  (NEW in v3 — what storytelling device
    #    does the video use to organize information?)
    # ══════════════════════════════════════════════════════════════════
    # Detect narrative structure from transcript + visual patterns.
    #
    # Options:
    #   - "Problem → Solution"          (problem framing → advice)
    #   - "Step-by-Step / How-To"       (sequential numbered steps)
    #   - "Analogy / Extended Metaphor"  (one metaphor carried throughout)
    #   - "Myth-Busting / Contrast"     (common belief → correction)
    #   - "Historical / Chronological"   (time-ordered narration)
    #   - "Listicle / Tips"             (numbered tips, moves, rules)
    #   - "Character Journey / Story Arc" (protagonist transformation)
    #   - "Data-Driven / Analytical"     (charts + formulas + evidence)
    #   - "Direct Explanation"           (straightforward concept definition)

    # Transcript pattern detectors
    _problem_solution_kw = [
        "here's how", "the problem", "the solution", "the fix",
        "what you can do", "move one", "move two", "step one",
        "step two", "here's what to do", "the trick is",
        "here's the secret",
    ]
    _howto_kw = [
        "step 1", "step 2", "step one", "step two", "first,",
        "second,", "third,", "how to", "tutorial", "guide",
    ]
    _mythbust_kw = [
        "most people", "you've been told", "actually", "the truth is",
        "myth", "misconception", "wrong", "not what you think",
        "without stereotypes", "isn't always",
    ]
    _listicle_kw = [
        "number one", "number two", "tip one", "tip two",
        "rule one", "rule two", "move one", "move two",
        "first thing", "second thing", "three things",
        "here are", "top 5", "top 10",
    ]
    _chronological_kw = [
        "in the 19th century", "in the 18th century", "in the 1800s",
        "during the war", "after the war", "by 1900", "medieval",
        "ancient", "colonial", "century", "historians", "era",
        "before dawn", "years later", "decades ago",
    ]
    _character_journey_kw = [
        "once upon", "our hero", "she discovered", "he learned",
        "his journey", "her journey", "overcame", "transformed",
        "grew up", "became", "started as",
    ]

    has_problem_solution = any(
        kw in full_transcript_lower for kw in _problem_solution_kw
    )
    has_howto = any(kw in full_transcript_lower for kw in _howto_kw)
    has_mythbust = any(kw in full_transcript_lower for kw in _mythbust_kw)
    has_listicle = any(kw in full_transcript_lower for kw in _listicle_kw)
    has_chronological = any(
        kw in full_transcript_lower for kw in _chronological_kw
    )
    has_character_journey = any(
        kw in full_transcript_lower for kw in _character_journey_kw
    )

    # Priority classification
    # Note: problem-solution evaluated BEFORE myth-busting because short
    # keywords like "wrong" in myth-busting produce false positives
    # (e.g. "timing goes wrong"), while problem-solution markers
    # ("here's how" + "move one"/"step one") are much more specific.
    if has_chronological and has_historical:
        narrative_structure = "Historical / Chronological"
    elif has_character_journey and character_strategy in (
        "Single Protagonist Arc", "Mascot-Led (non-human recurring character)"
    ):
        narrative_structure = "Character Journey / Story Arc"
    elif r_data_pres >= 0.40 and r_symbolism < 0.15:
        narrative_structure = "Data-Driven / Analytical"
    elif has_problem_solution and has_listicle:
        narrative_structure = "Problem → Solution"
    elif has_problem_solution:
        narrative_structure = "Problem → Solution"
    elif has_mythbust:
        narrative_structure = "Myth-Busting / Contrast"
    elif has_listicle:
        narrative_structure = "Listicle / Tips"
    elif has_howto:
        narrative_structure = "Step-by-Step / How-To"
    elif r_analogy >= 0.20 and r_fictional >= 0.20:
        narrative_structure = "Analogy / Extended Metaphor"
    elif r_symbolism >= 0.25:
        narrative_structure = "Analogy / Extended Metaphor"
    elif r_data_pres >= 0.25:
        narrative_structure = "Data-Driven / Analytical"
    else:
        narrative_structure = "Direct Explanation"

    # ══════════════════════════════════════════════════════════════════
    # 5) Visual Abstraction Index  (NEW in v3 — 1-5 graduated scale)
    # ══════════════════════════════════════════════════════════════════
    # 1 = Photorealistic / live-action
    # 2 = Stylized realism (3D Pixar-style, detailed 2D illustration)
    # 3 = Moderate abstraction (flat 2D cartoon, simplified environments)
    # 4 = High abstraction (line art, stick figures, simple shapes)
    # 5 = Maximum abstraction (floating objects on void, pure data/text)
    #
    # Computed as weighted score from feature ratios.
    _abs_score = 0.0
    _abs_weights = 0.0

    # Pull toward 1 (photorealistic)
    if r_photoreal > 0:
        _abs_score += 1.0 * r_photoreal * 3.0
        _abs_weights += r_photoreal * 3.0

    # Pull toward 2 (stylized realism)
    if r_3d > 0 and r_photoreal < 0.15:
        _abs_score += 2.0 * r_3d * 2.5
        _abs_weights += r_3d * 2.5
    if r_textured > 0:
        _abs_score += 2.5 * r_textured * 2.0
        _abs_weights += r_textured * 2.0
    if r_real_env > 0 and r_stylized > 0:
        _w = min(r_real_env, r_stylized) * 1.5
        _abs_score += 2.0 * _w
        _abs_weights += _w

    # Pull toward 3 (moderate abstraction)
    if r_2d_flat > 0:
        _abs_score += 3.0 * r_2d_flat * 2.0
        _abs_weights += r_2d_flat * 2.0
    if r_2d_broad > 0 and r_lineart < 0.20:
        _abs_score += 3.0 * r_2d_broad * 1.0
        _abs_weights += r_2d_broad * 1.0

    # Pull toward 4 (high abstraction)
    if r_lineart > 0:
        _abs_score += 4.0 * r_lineart * 3.0
        _abs_weights += r_lineart * 3.0

    # Pull toward 5 (maximum abstraction)
    if r_abstract_bg > 0 and r_real_env < 0.15:
        _w5 = r_abstract_bg * 2.0
        _abs_score += 5.0 * _w5
        _abs_weights += _w5
    if r_data_pres > 0.30 and r_human_char < 0.20:
        _w5d = r_data_pres * 1.5
        _abs_score += 5.0 * _w5d
        _abs_weights += _w5d

    # Compute final index (default to 3 if no signals)
    if _abs_weights > 0:
        _raw_index = _abs_score / _abs_weights
    else:
        _raw_index = 3.0

    # Round to nearest integer, clamp 1-5
    abstraction_index = max(1, min(5, round(_raw_index)))

    # Map to descriptive label
    _abstraction_labels = {
        1: "1 — Photorealistic",
        2: "2 — Stylized Realism",
        3: "3 — Moderate Abstraction",
        4: "4 — High Abstraction",
        5: "5 — Maximum Abstraction (icons/text/void)",
    }
    visual_abstraction_index = _abstraction_labels[abstraction_index]

    # ══════════════════════════════════════════════════════════════════
    # 6) Visual Density
    # ══════════════════════════════════════════════════════════════════
    # Measures how many visual layers compete for attention per shot.
    # Uses text overlays, data elements, UI widgets, and scene complexity.

    # Extra graphical elements beyond the primary subject
    _extra_elements_kw = [
        "chart", "icon", "ui", "logo", "badge", "seal", "phone",
        "app", "button", "tab", "graph", "infographic", "diagram",
        "dashboard", "browser",
    ]
    r_extra = _shot_ratio(_extra_elements_kw)

    # Plain / minimal backgrounds
    _plain_bg_kw = [
        "solid background", "plain background", "white background",
        "clean background", "simple background", "solid color",
        "black void",
    ]
    r_plain = _shot_ratio(_plain_bg_kw)

    # Layered complexity signals
    _layered_kw = [
        "split screen", "multiple", "collage", "grid", "overlay",
        "composite", "layered", "picture-in-picture",
    ]
    r_layered = _shot_ratio(_layered_kw)

    if r_text_overlay >= 0.30 and r_extra >= 0.20:
        visual_density = "High (layered text + icons + backgrounds)"
    elif r_layered >= 0.20 or (r_text_overlay >= 0.20 and r_extra >= 0.30):
        visual_density = "High (layered text + icons + backgrounds)"
    elif r_plain >= 0.60 and r_extra < 0.15:
        visual_density = "Minimal (single subject, clean background)"
    elif r_plain >= 0.40 and r_text_overlay < 0.15:
        visual_density = "Sparse"
    else:
        visual_density = "Moderate"

    # ══════════════════════════════════════════════════════════════════
    # 7) Camera / Editing Language
    # ══════════════════════════════════════════════════════════════════

    if is_vertical and avg_shot_len <= 4.0 and has_cta:
        camera_editing = "Social Vertical Punch (fast cuts, kinetic text, CTA)"
    elif is_vertical and is_short_form and has_cta:
        camera_editing = "Social Vertical Punch (fast cuts, kinetic text, CTA)"
    elif r_cinematic >= 0.20:
        camera_editing = "Cinematic (depth-of-field, dramatic lighting, camera motion)"
    elif r_institutional >= 0.15 and r_data_pres >= 0.15 and r_text_overlay >= 0.20:
        camera_editing = "Presentation Deck (structured title cards + charts + logos)"
    elif is_vertical and avg_shot_len <= 4.0:
        camera_editing = "Social Vertical Punch (fast cuts, kinetic text, CTA)"
    elif avg_shot_len >= 8.0 and r_motion < 0.10:
        camera_editing = "Slow / Deliberate (long holds, minimal cuts)"
    elif r_motion >= 0.25:
        camera_editing = "Dynamic (frequent camera motion, zooms, tracking)"
    else:
        # Static vs Mixed
        variety = len(set(v[:40] for v in visual_list))
        if r_motion < 0.10 and variety < total * 0.5:
            camera_editing = "Static Slides / Centered Frames"
        else:
            camera_editing = "Mixed / Moderate Pacing"

    # ══════════════════════════════════════════════════════════════════
    # 8) Tonal Positioning  (strict hierarchy — stop on first match)
    # ══════════════════════════════════════════════════════════════════

    # Pre-compute child content guard (used in Gen Z vs Child priority)
    _is_child_content = (r_mascot >= 0.40 or r_child_vis >= 0.25) and grade_num <= 5

    tonal_positioning = None

    # ── PRIORITY 1: Institutional / Formal ──
    if duration > 150 and is_horizontal:
        tonal_positioning = "Institutional / Formal"
    elif r_institutional >= 0.20:
        tonal_positioning = "Institutional / Formal"
    elif has_historical:
        tonal_positioning = "Institutional / Formal"
    elif (not has_cta and not has_genz_slang and grade_num >= 9
          and r_corporate < 0.25 and r_dark_vis < 0.20):
        tonal_positioning = "Institutional / Formal"

    # ── PRIORITY 2: Corporate Professional ──
    if tonal_positioning is None:
        if r_corporate >= 0.40 and not has_cta and not has_genz_slang:
            tonal_positioning = "Corporate Professional"

    # ── PRIORITY 3: Gen Z Social / Snappy ──
    #    (skipped if mascot-heavy + low grade → child content)
    if tonal_positioning is None and not _is_child_content:
        if is_vertical and avg_shot_len <= 4.0:
            tonal_positioning = "Gen Z Social / Snappy"
        elif has_cta and has_genz_slang:
            tonal_positioning = "Gen Z Social / Snappy"
        elif duration <= 90 and is_vertical and has_cta:
            tonal_positioning = "Gen Z Social / Snappy"
        elif has_genz_slang:
            tonal_positioning = "Gen Z Social / Snappy"

    # ── PRIORITY 4: Child-Friendly / Whimsical ──
    if tonal_positioning is None:
        if _is_child_content:
            tonal_positioning = "Child-Friendly / Whimsical"

    # ── PRIORITY 5: Dark Editorial / Satirical ──
    if tonal_positioning is None:
        _is_symbolic_narrative = narrative_structure in (
            "Analogy / Extended Metaphor",
        ) or r_symbolism >= 0.25
        if _is_symbolic_narrative and r_dark_vis >= 0.30 and has_moral_framing:
            tonal_positioning = "Dark Editorial / Satirical"
        elif r_dark_vis >= 0.30 and has_moral_framing:
            tonal_positioning = "Dark Editorial / Satirical"

    # ── DEFAULT FALLBACK ──
    if tonal_positioning is None:
        tonal_positioning = "Corporate Professional"

    # ================================================================
    # Assemble output  (v3 canonical keys + v2 backward-compat aliases)
    # ================================================================
    return {
        # ── v3 canonical keys ──
        "rendering_class": rendering_class,
        "world_type": world_type,
        "character_strategy": character_strategy,
        "narrative_structure": narrative_structure,
        "visual_abstraction_index": visual_abstraction_index,
        "visual_density": visual_density,
        "camera_editing": camera_editing,
        "tonal_positioning": tonal_positioning,
        # ── v2 backward-compat aliases (to be removed in future) ──
        "metaphor_mode": narrative_structure,
    }


# =============================================================================
# Alignment: Match Transcripts to Shots
# =============================================================================

def align_transcript_to_shots(
    shots: list[Shot],
    transcript: list[TranscriptSegment]
) -> list[Shot]:
    """
    Align transcribed audio segments to visual shots based on timestamps.

    Each shot gets the dialogue that overlaps with its time range.
    """
    for shot in shots:
        # Find all transcript segments that overlap with this shot
        overlapping_segments = []

        for segment in transcript:
            # Check for overlap
            if segment.end_time > shot.start_time and segment.start_time < shot.end_time:
                overlapping_segments.append(segment)

        # Combine overlapping segments into dialogue
        if overlapping_segments:
            dialogue_parts = []
            for seg in overlapping_segments:
                speaker = infer_speaker(seg.text)
                dialogue_parts.append(f'{speaker}: "{seg.text}"')
            shot.dialogue = " ".join(dialogue_parts)
        else:
            shot.dialogue = "(no dialogue)"

    return shots


# =============================================================================
# Output Generation
# =============================================================================

def generate_markdown_output(
    shots: list[Shot],
    metadata: VideoMetadata,
    frames_dir: str,
    visual_style: Optional[dict] = None
) -> str:
    """Generate markdown output with embedded frames and rich header."""
    # Build markdown
    lines = [
        f"# Video Analysis: {metadata.title}",
        "",
        "## Video Overview",
        "",
        "| Property | Value |",
        "|----------|-------|",
        f"| **Runtime** | {metadata.duration_str} |",
        f"| **Format** | {metadata.aspect_ratio} |",
        f"| **Word Count** | {metadata.word_count} |",
        f"| **Image Count** | {metadata.image_count} |",
        f"| **Reading Level** | {metadata.reading_level or 'N/A'} |",
        f"| **Story Structure** | {metadata.story_structure or 'N/A'} |",
        "",
    ]

    # Add Stylistic Fingerprint (v3) section (between Video Overview and Key Details)
    if visual_style:
        lines.extend([
            "## Stylistic Fingerprint (v3)",
            "",
            "| Pillar | Value |",
            "|--------|-------|",
            f"| **Rendering Class** | {visual_style['rendering_class']} |",
            f"| **World Type** | {visual_style['world_type']} |",
            f"| **Character Strategy** | {visual_style['character_strategy']} |",
            f"| **Narrative Structure** | {visual_style.get('narrative_structure', 'N/A')} |",
            f"| **Visual Abstraction Index** | {visual_style.get('visual_abstraction_index', 'N/A')} |",
            f"| **Visual Density** | {visual_style['visual_density']} |",
            f"| **Camera / Editing Language** | {visual_style['camera_editing']} |",
            f"| **Tonal Positioning** | {visual_style['tonal_positioning']} |",
            "",
        ])

    # Add key details if available
    if metadata.key_details:
        lines.extend([
            "### Key Details",
            "",
            metadata.key_details,
            "",
        ])

    # Add storyboard section
    lines.extend([
        "## Shot-by-Shot Breakdown",
        "",
        "| Shot | Timestamp | Frame | Visual | Voiceover/Dialogue |",
        "|------|-----------|-------|--------|-------------------|"
    ])

    for shot in shots:
        # Escape pipe characters in content
        visual = shot.visual_description.replace("|", "\\|")
        dialogue = shot.dialogue.replace("|", "\\|")

        # Create relative path to frame image
        if shot.frame_path and os.path.exists(shot.frame_path):
            frame_filename = os.path.basename(shot.frame_path)
            frame_ref = f"![Shot {shot.number}](../frames/{frame_filename})"
        else:
            frame_ref = "[no frame]"

        lines.append(
            f"| {shot.number} | {shot.timestamp_str} | {frame_ref} | {visual} | {dialogue} |"
        )

    return "\n".join(lines)


def generate_docx_output(
    shots: list[Shot],
    metadata: VideoMetadata,
    output_path: str,
    visual_style: Optional[dict] = None
) -> str:
    """Generate Word document (.docx) output with table and embedded frames."""
    # Lazy import to avoid import issues
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required for Word document output. Install with: pip install python-docx")

    # Create document
    doc = Document()

    # Add title
    doc.add_heading(f"Video Analysis: {metadata.title}", 0)

    # Add rich metadata header
    doc.add_heading("Video Overview", level=1)

    # Create metadata table for clean formatting
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.style = "Table Grid"

    meta_items = [
        ("Runtime", metadata.duration_str),
        ("Format", metadata.aspect_ratio),
        ("Word Count", str(metadata.word_count)),
        ("Image Count", str(metadata.image_count)),
        ("Reading Level", metadata.reading_level or "N/A"),
        ("Story Structure", metadata.story_structure or "N/A"),
    ]

    for i, (label, value) in enumerate(meta_items):
        row = meta_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        # Bold the labels
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Set column widths for metadata table
    for row in meta_table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)

    doc.add_paragraph()  # Spacer

    # Add Stylistic Fingerprint (v3) section (between Video Overview and Key Details)
    if visual_style:
        doc.add_heading("Stylistic Fingerprint (v3)", level=1)

        fp_items = [
            ("Rendering Class", visual_style["rendering_class"]),
            ("World Type", visual_style["world_type"]),
            ("Character Strategy", visual_style["character_strategy"]),
            ("Narrative Structure", visual_style.get("narrative_structure", "N/A")),
            ("Visual Abstraction Index", visual_style.get("visual_abstraction_index", "N/A")),
            ("Visual Density", visual_style["visual_density"]),
            ("Camera / Editing Language", visual_style["camera_editing"]),
            ("Tonal Positioning", visual_style["tonal_positioning"]),
        ]

        fp_table = doc.add_table(rows=len(fp_items), cols=2)
        fp_table.style = "Table Grid"

        for i, (label, value) in enumerate(fp_items):
            row = fp_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for row in fp_table.rows:
            row.cells[0].width = Inches(2.5)
            row.cells[1].width = Inches(4.0)

        doc.add_paragraph()  # Spacer

    # Add key details if available
    if metadata.key_details:
        doc.add_heading("Key Details", level=2)
        doc.add_paragraph(metadata.key_details)

    doc.add_paragraph()  # Spacer

    # Add storyboard section
    doc.add_heading("Shot-by-Shot Breakdown", level=1)

    # Create table with 5 columns: Shot, Timestamp, Frame, Visual, Dialogue
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Shot"
    header_cells[1].text = "Timestamp"
    header_cells[2].text = "Frame"
    header_cells[3].text = "Visual"
    header_cells[4].text = "Voiceover/Dialogue"

    # Make headers bold
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Determine frame width based on orientation
    if metadata.orientation == "vertical":
        frame_width = Inches(0.8)  # Narrower for vertical
    else:
        frame_width = Inches(1.2)  # Wider for horizontal

    # Add data rows
    for shot in shots:
        row_cells = table.add_row().cells
        row_cells[0].text = str(shot.number)
        row_cells[1].text = shot.timestamp_str

        # Add frame image if available
        if shot.frame_path and os.path.exists(shot.frame_path):
            paragraph = row_cells[2].paragraphs[0]
            run = paragraph.add_run()
            try:
                run.add_picture(shot.frame_path, width=frame_width)
            except Exception:
                row_cells[2].text = "[image]"
        else:
            row_cells[2].text = "[no frame]"

        row_cells[3].text = shot.visual_description
        row_cells[4].text = shot.dialogue

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(0.4)   # Shot
        row.cells[1].width = Inches(0.8)   # Timestamp
        row.cells[2].width = Inches(1.3)   # Frame
        row.cells[3].width = Inches(2.5)   # Visual
        row.cells[4].width = Inches(2.5)   # Dialogue

    # Save document
    doc.save(output_path)
    return output_path


# =============================================================================
# Main Processing Pipeline
# =============================================================================

def process_video(
    input_path: str,
    interval: Optional[float] = None,
    output_format: str = "docx",
    cookies_file: Optional[str] = None,
    verbose: bool = False
) -> str:
    """
    Main processing pipeline for video analysis.

    Args:
        input_path: Path to video file or URL
        interval: If set, use fixed interval instead of scene detection
        output_format: Output format - "docx" or "md"
        cookies_file: Path to cookies.txt for authenticated downloads
        verbose: Print progress information

    Returns:
        Path to generated output file
    """
    # Get script directory for output folder
    script_dir = Path(__file__).parent.absolute()
    output_base = script_dir / "output"

    # Create temp directory for processing
    with tempfile.TemporaryDirectory() as temp_dir:
        # Handle URL input
        if is_url(input_path):
            if verbose:
                print("=" * 50)
                print("Step 1: Downloading video")
                print("=" * 50)
            video_path, title = download_video(input_path, temp_dir, cookies_file, verbose)
        else:
            video_path = input_path
            title = get_video_title(input_path)

        # Create output folder structure: output/[video name]/frames and output/[video name]/storyboard
        safe_title = sanitize_filename(title)
        video_output_dir = output_base / safe_title
        frames_dir = video_output_dir / "frames"
        storyboard_dir = video_output_dir / "storyboard"

        os.makedirs(frames_dir, exist_ok=True)
        os.makedirs(storyboard_dir, exist_ok=True)

        if verbose:
            print(f"Output folder: {video_output_dir}")

        # Get video duration and dimensions
        duration = get_video_duration(video_path)
        width, height, orientation = get_video_dimensions(video_path)
        if verbose:
            print(f"Video duration: {duration:.1f} seconds")
            print(f"Video dimensions: {width}x{height} ({orientation})")

        # Step 2: Extract frames / detect scenes
        if verbose:
            print("=" * 50)
            print("Step 2: Frame extraction / scene detection")
            print("=" * 50)

        if interval:
            # Use fixed interval extraction
            frames = extract_frames_at_interval(
                video_path, str(frames_dir), interval, verbose
            )
            shots = create_shots_from_interval(frames, duration, interval)
        else:
            # Use scene detection
            scenes = detect_scenes(video_path, verbose)

            # Fallback to interval if no scenes detected
            if not scenes:
                if verbose:
                    print("No scenes detected, falling back to 3-second intervals")
                frames = extract_frames_at_interval(
                    video_path, str(frames_dir), 3.0, verbose
                )
                shots = create_shots_from_interval(frames, duration, 3.0)
            else:
                shots = create_shots_from_scenes(
                    scenes, video_path, str(frames_dir), verbose
                )

        if verbose:
            print(f"Created {len(shots)} shots")

        # Step 3: Transcribe audio
        if verbose:
            print("=" * 50)
            print("Step 3: Audio transcription")
            print("=" * 50)

        transcript = transcribe_audio(video_path, verbose=verbose)

        # Step 4: Analyze visuals with Claude
        if verbose:
            print("=" * 50)
            print("Step 4: Visual analysis with Claude")
            print("=" * 50)

        shots = analyze_all_frames(shots, orientation, verbose)

        # Step 5: Align transcript to shots
        if verbose:
            print("=" * 50)
            print("Step 5: Aligning transcript to shots")
            print("=" * 50)

        shots = align_transcript_to_shots(shots, transcript)

        # Step 5b: Analyze content for metadata (before fingerprint, since
        #          fingerprint v2 uses reading_level for tonal positioning)
        if verbose:
            print("=" * 50)
            print("Step 5b: Content analysis for metadata")
            print("=" * 50)

        key_details, reading_level, story_structure = analyze_video_content(
            transcript, shots, title, verbose
        )

        # Step 6: Stylistic fingerprint (v3)
        if verbose:
            print("=" * 50)
            print("Step 6: Stylistic Fingerprint v3 (production-grade)")
            print("=" * 50)

        visual_style = analyze_stylistic_fingerprint(
            shots, transcript, duration,
            orientation=orientation,
            reading_level=reading_level,
        )

        if verbose:
            print(f"  Rendering Class: {visual_style['rendering_class']}")
            print(f"  World Type: {visual_style['world_type']}")
            print(f"  Character Strategy: {visual_style['character_strategy']}")
            print(f"  Narrative Structure: {visual_style['narrative_structure']}")
            print(f"  Visual Abstraction: {visual_style['visual_abstraction_index']}")
            print(f"  Visual Density: {visual_style['visual_density']}")
            print(f"  Camera/Editing: {visual_style['camera_editing']}")
            print(f"  Tonal Positioning: {visual_style['tonal_positioning']}")

        # Calculate word count from transcript
        full_transcript = " ".join([seg.text for seg in transcript])
        word_count = len(full_transcript.split())

        # Create metadata object
        metadata = VideoMetadata(
            title=title,
            duration=duration,
            width=width,
            height=height,
            orientation=orientation,
            word_count=word_count,
            image_count=len(shots),
            key_details=key_details,
            reading_level=reading_level,
            story_structure=story_structure
        )

        # Step 7: Generate output
        if verbose:
            print("=" * 50)
            print("Step 7: Generating output")
            print("=" * 50)

        # Generate outputs with video title prefix
        if output_format.lower() == "docx":
            # Create both generic and titled versions
            generic_path = storyboard_dir / "breakdown.docx"
            titled_path = storyboard_dir / f"{safe_title}_breakdown.docx"
            generate_docx_output(shots, metadata, str(generic_path), visual_style)
            generate_docx_output(shots, metadata, str(titled_path), visual_style)
            output_path = titled_path
        else:
            generic_path = storyboard_dir / "breakdown.md"
            titled_path = storyboard_dir / f"{safe_title}_breakdown.md"
            markdown = generate_markdown_output(shots, metadata, str(frames_dir), visual_style)
            with open(generic_path, "w") as f:
                f.write(markdown)
            with open(titled_path, "w") as f:
                f.write(markdown)
            output_path = titled_path

        if verbose:
            print(f"\nOutputs written to:")
            print(f"  - {generic_path}")
            print(f"  - {titled_path}")
            print(f"Frames saved to: {frames_dir}")

        return str(output_path)


# =============================================================================
# CLI Entry Point
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Analyze educational videos and output shot-by-shot breakdown tables.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python analyze_video.py "https://youtube.com/watch?v=xxx" --verbose
  python analyze_video.py input.mp4 --verbose
  python analyze_video.py input.mp4 --format md --interval 3

Output structure:
  output/
  └── [video name]/
      ├── frames/       (extracted frames)
      └── storyboard/   (breakdown.docx or breakdown.md)

Environment:
  ANTHROPIC_API_KEY - Required for Claude Vision analysis
        """
    )

    parser.add_argument(
        "input",
        help="Input video file path or YouTube/Vimeo URL"
    )

    parser.add_argument(
        "--format", "-f",
        choices=["docx", "md"],
        default="docx",
        help="Output format: docx (Word) or md (Markdown). Default: docx"
    )

    parser.add_argument(
        "--interval",
        type=float,
        default=None,
        help="Extract frames every N seconds instead of using scene detection"
    )

    parser.add_argument(
        "--verbose", "-v",
        action="store_true",
        help="Show progress during processing"
    )

    parser.add_argument(
        "--cookies", "-c",
        type=str,
        default=None,
        help="Path to cookies.txt file for sites requiring authentication (TikTok, etc.)"
    )

    args = parser.parse_args()

    # Check dependencies
    missing = check_dependencies()
    if missing:
        print("Missing dependencies:")
        for dep in missing:
            print(f"  - {dep}")
        print("\nInstall with:")
        print("  pip install openai-whisper scenedetect[opencv] anthropic Pillow")
        print("  brew install ffmpeg yt-dlp  # or apt-get install on Linux")
        sys.exit(1)

    # Check for API key
    if not os.environ.get("ANTHROPIC_API_KEY"):
        print("Error: ANTHROPIC_API_KEY environment variable is required")
        print("Set it with: export ANTHROPIC_API_KEY='your-key-here'")
        sys.exit(1)

    # Validate input
    if not is_url(args.input) and not os.path.exists(args.input):
        print(f"Error: Input file not found: {args.input}")
        sys.exit(1)

    try:
        output_path = process_video(
            args.input,
            interval=args.interval,
            output_format=args.format,
            cookies_file=args.cookies,
            verbose=args.verbose
        )
        print(f"Analysis complete! Output saved to: {output_path}")
    except KeyboardInterrupt:
        print("\nInterrupted by user")
        sys.exit(130)
    except Exception as e:
        print(f"Error: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()
